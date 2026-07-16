"""
File extraction endpoint.

Accepts one or more uploaded files — plain text/code, PDFs, Word docs,
Excel sheets, or ZIP archives (which are unpacked and each member
extracted in turn) — and returns their text content so the frontend can
fold it into a chat message for the model to read.

Nothing is written to disk or persisted to the database; extraction
happens entirely in memory for the duration of the request.
"""
from __future__ import annotations

import io
import zipfile

from fastapi import APIRouter, Depends, File, UploadFile
from pydantic import BaseModel

from app.core.deps import get_current_user
from app.db.models import User

router = APIRouter(prefix="/api/files", tags=["files"])

MAX_FILE_BYTES = 20 * 1024 * 1024  # 20 MB per individual uploaded/zipped file
MAX_ZIP_MEMBERS = 50  # don't unpack absurdly large archives
MAX_EXTRACTED_CHARS = 200_000  # per-file cap so one huge doc can't blow the model's context

TEXT_EXTENSIONS = {
    "txt", "md", "markdown", "json", "csv", "tsv", "log", "yaml", "yml",
    "xml", "html", "htm", "css", "scss", "js", "jsx", "ts", "tsx", "py",
    "java", "c", "h", "cpp", "hpp", "cs", "go", "rs", "rb", "php", "sh",
    "bash", "sql", "ini", "toml", "cfg", "conf", "env", "r", "kt", "swift",
    "vue", "svelte", "gitignore",
}


class ExtractedFile(BaseModel):
    name: str
    text: str
    truncated: bool = False
    warning: str | None = None


class ExtractResponse(BaseModel):
    files: list[ExtractedFile]


def _truncate(text: str) -> tuple[str, bool]:
    if len(text) > MAX_EXTRACTED_CHARS:
        return text[:MAX_EXTRACTED_CHARS], True
    return text, False


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pages.append("")
    return "\n\n".join(pages).strip()


def _extract_docx(data: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts).strip()


def _extract_xlsx(data: bytes) -> str:
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    lines: list[str] = []
    for sheet in wb.worksheets:
        lines.append(f"# Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            lines.append(", ".join("" if c is None else str(c) for c in row))
    return "\n".join(lines).strip()


def _extract_plaintext(data: bytes) -> str:
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    raise ValueError("could not decode as text")


def _looks_binary(text: str) -> bool:
    if "\u0000" in text:
        return True
    replacement = text.count("\ufffd")
    return replacement > max(5, len(text) * 0.01)


def _extract_single(name: str, data: bytes) -> ExtractedFile:
    if len(data) > MAX_FILE_BYTES:
        return ExtractedFile(
            name=name,
            text="",
            warning=f"Skipped — larger than the {MAX_FILE_BYTES // (1024 * 1024)} MB per-file limit.",
        )

    ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""

    try:
        if ext == "pdf":
            text = _extract_pdf(data)
        elif ext == "docx":
            text = _extract_docx(data)
        elif ext in ("xlsx", "xlsm"):
            text = _extract_xlsx(data)
        elif ext in TEXT_EXTENSIONS or ext == "":
            text = _extract_plaintext(data)
            if _looks_binary(text):
                return ExtractedFile(name=name, text="", warning="Appears to be a binary file — couldn't read it as text.")
        elif ext == "doc":
            return ExtractedFile(name=name, text="", warning="Old-style .doc isn't supported — please save as .docx and re-attach.")
        elif ext in ("xls",):
            return ExtractedFile(name=name, text="", warning="Old-style .xls isn't supported — please save as .xlsx and re-attach.")
        else:
            return ExtractedFile(name=name, text="", warning=f"Unsupported file type '.{ext}' — couldn't extract text.")
    except Exception as exc:  # noqa: BLE001 - surface any extraction failure as a warning, not a 500
        return ExtractedFile(name=name, text="", warning=f"Failed to read this file: {exc}")

    text = text.strip()
    if not text:
        return ExtractedFile(
            name=name,
            text="",
            warning="No extractable text found (it may be a scanned/image-only document).",
        )

    text, truncated = _truncate(text)
    return ExtractedFile(name=name, text=text, truncated=truncated)


def _extract_zip(name: str, data: bytes) -> list[ExtractedFile]:
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            members = [m for m in zf.infolist() if not m.is_dir()]
            results: list[ExtractedFile] = []
            for member in members[:MAX_ZIP_MEMBERS]:
                try:
                    member_bytes = zf.read(member)
                except Exception as exc:  # noqa: BLE001
                    results.append(ExtractedFile(name=f"{name}/{member.filename}", text="", warning=f"Couldn't extract from archive: {exc}"))
                    continue
                extracted = _extract_single(member.filename, member_bytes)
                extracted.name = f"{name}/{member.filename}"
                results.append(extracted)
            if len(members) > MAX_ZIP_MEMBERS:
                results.append(
                    ExtractedFile(
                        name=name,
                        text="",
                        warning=f"Archive has more than {MAX_ZIP_MEMBERS} files — only the first {MAX_ZIP_MEMBERS} were read.",
                    )
                )
            return results
    except zipfile.BadZipFile:
        return [ExtractedFile(name=name, text="", warning="Not a valid ZIP archive.")]


@router.post("/extract", response_model=ExtractResponse)
async def extract_files(
    uploads: list[UploadFile] = File(...),
    _user: User = Depends(get_current_user),
) -> ExtractResponse:
    results: list[ExtractedFile] = []
    for upload in uploads:
        name = upload.filename or "file"
        data = await upload.read()

        if name.lower().endswith(".zip"):
            results.extend(_extract_zip(name, data))
        else:
            results.append(_extract_single(name, data))

    return ExtractResponse(files=results)